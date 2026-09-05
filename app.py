```python
import io
import json
import os
import time

import streamlit as st
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document


# =========================================================
# STREAMLIT CONFIG
# =========================================================

st.set_page_config(
    page_title="Resume ATS Analyzer",
    page_icon="📄",
    layout="wide",
)


# =========================================================
# GEMINI CONFIG
# =========================================================

MODEL_NAME = "gemini-3.8-flash"

MAX_RETRIES = 4
INITIAL_RETRY_DELAY = 5


# =========================================================
# API KEY
# =========================================================

def get_api_key():
    """Read Gemini API key from Streamlit Secrets or environment variable."""

    try:
        secret_key = st.secrets.get("GEMINI_API_KEY")
    except Exception:
        secret_key = None

    return secret_key or os.getenv("GEMINI_API_KEY")


# =========================================================
# RESUME TEXT EXTRACTION
# =========================================================

def extract_resume_text(uploaded_file):
    """Extract text from PDF, DOCX, or TXT files."""

    file_bytes = uploaded_file.getvalue()
    file_type = uploaded_file.type or ""
    file_name = uploaded_file.name.lower()

    # -------------------------
    # PDF
    # -------------------------

    if file_type == "application/pdf" or file_name.endswith(".pdf"):

        reader = PdfReader(io.BytesIO(file_bytes))

        pages = [
            page.extract_text() or ""
            for page in reader.pages
        ]

        text = "\n".join(pages)

    # -------------------------
    # DOCX
    # -------------------------

    elif (
        file_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or file_name.endswith(".docx")
    ):

        document = Document(io.BytesIO(file_bytes))

        parts = [
            p.text
            for p in document.paragraphs
            if p.text.strip()
        ]

        # Also read tables
        for table in document.tables:

            for row in table.rows:

                parts.append(
                    " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                    )
                )

        text = "\n".join(parts)

    # -------------------------
    # TXT / MD
    # -------------------------

    elif file_name.endswith((".txt", ".md")) or file_type.startswith("text/"):

        text = file_bytes.decode(
            "utf-8",
            errors="ignore"
        )

    else:

        raise ValueError(
            "Unsupported file type. "
            "Please upload PDF, DOCX, or TXT."
        )

    # Clean text
    text = "\n".join(
        line.rstrip()
        for line in text.splitlines()
    ).strip()

    if not text:

        raise ValueError(
            "No readable text was found. "
            "If your PDF is a scanned image, "
            "please use a text-based PDF or DOCX."
        )

    return text


# =========================================================
# GEMINI REQUEST
# =========================================================

def generate_gemini_response(client, prompt, schema):
    """
    Send request to Gemini with retry handling.

    Handles temporary 503 / 429 errors using
    exponential backoff.
    """

    last_error = None

    for attempt in range(MAX_RETRIES):

        try:

            response = client.models.generate_content(
                model=MODEL_NAME,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.2,
                    response_mime_type="application/json",
                    response_schema=schema,
                ),
            )

            return response

        except Exception as exc:

            last_error = exc

            error_text = str(exc)

            # Retry only temporary server/rate-limit problems
            is_retryable = (
                "503" in error_text
                or "UNAVAILABLE" in error_text
                or "429" in error_text
                or "RESOURCE_EXHAUSTED" in error_text
            )

            if not is_retryable:

                raise

            # Last attempt
            if attempt == MAX_RETRIES - 1:

                raise RuntimeError(
                    "Gemini is currently unavailable after "
                    f"{MAX_RETRIES} attempts.\n\n"
                    f"Original error: {error_text}"
                )

            delay = INITIAL_RETRY_DELAY * (2 ** attempt)

            st.warning(
                f"Gemini server is temporarily busy. "
                f"Retrying in {delay} seconds..."
            )

            time.sleep(delay)

    raise RuntimeError(
        f"Gemini request failed: {last_error}"
    )


# =========================================================
# RESUME ANALYSIS
# =========================================================

def analyze_resume(resume_text, job_description, api_key):
    """Ask Gemini for a structured ATS-style analysis."""

    client = genai.Client(
        api_key=api_key
    )

    # -----------------------------------------------------
    # JSON SCHEMA
    # -----------------------------------------------------

    schema = {
        "type": "object",

        "properties": {

            "ats_score": {
                "type": "integer",
                "description": (
                    "Overall ATS readiness score from 0 to 100."
                ),
            },

            "summary": {
                "type": "string",
                "description": (
                    "Short overall assessment of the resume."
                ),
            },

            "strengths": {
                "type": "array",

                "items": {
                    "type": "string"
                },

                "description": (
                    "Strong parts of the resume."
                ),
            },

            "improvements": {

                "type": "array",

                "items": {

                    "type": "object",

                    "properties": {

                        "issue": {
                            "type": "string"
                        },

                        "recommendation": {
                            "type": "string"
                        },

                        "priority": {

                            "type": "string",

                            "enum": [
                                "High",
                                "Medium",
                                "Low",
                            ],
                        },
                    },

                    "required": [
                        "issue",
                        "recommendation",
                        "priority",
                    ],
                },

                "description": (
                    "Specific improvements the candidate should make."
                ),
            },

            "keywords_missing": {

                "type": "array",

                "items": {
                    "type": "string"
                },

                "description": (
                    "Important job-related keywords "
                    "missing or weakly represented."
                ),
            },

            "ats_checks": {

                "type": "array",

                "items": {

                    "type": "object",

                    "properties": {

                        "check": {
                            "type": "string"
                        },

                        "status": {

                            "type": "string",

                            "enum": [
                                "Good",
                                "Needs improvement",
                                "Risk",
                            ],
                        },

                        "explanation": {
                            "type": "string"
                        },
                    },

                    "required": [
                        "check",
                        "status",
                        "explanation",
                    ],
                },

                "description": (
                    "Checks for formatting, sections, "
                    "keywords, readability, and measurable achievements."
                ),
            },
        },

        "required": [
            "ats_score",
            "summary",
            "strengths",
            "improvements",
            "keywords_missing",
            "ats_checks",
        ],
    }

    # -----------------------------------------------------
    # JOB DESCRIPTION
    # -----------------------------------------------------

    if job_description and job_description.strip():

        job_context = job_description.strip()

    else:

        job_context = (
            "No specific job description was provided. "
            "Evaluate general ATS readiness."
        )

    # -----------------------------------------------------
    # PROMPT
    # -----------------------------------------------------

    prompt = f"""
You are an expert resume and ATS analyst.

Analyze the resume below.

Produce:

1. An ATS-readiness score from 0 to 100.
2. A short overall summary.
3. Resume strengths.
4. Specific improvements.
5. Missing or weak keywords.
6. ATS formatting and readability checks.

Important rules:

- This is an ATS-readiness estimate.
- It is NOT a score produced by a real ATS vendor.
- Do not invent qualifications.
- Do not invent jobs.
- Do not invent degrees.
- Do not invent dates.
- Do not invent achievements.
- Base every recommendation only on the supplied resume
  and job description.
- Check whether sections are clear and conventional.
- Check keyword relevance.
- Check measurable achievements.
- Check action verbs.
- Check clarity and consistency.
- Check readability.
- Check possible ATS parsing problems.
- If a job description is supplied, compare the resume against it.
- Identify important missing or weak keywords.
- Keep recommendations practical and concise.

JOB DESCRIPTION:

{job_context}


RESUME:

{resume_text}
"""

    # -----------------------------------------------------
    # CALL GEMINI
    # -----------------------------------------------------

    response = generate_gemini_response(
        client=client,
        prompt=prompt,
        schema=schema,
    )

    if not response.text:

        raise RuntimeError(
            "Gemini returned an empty response."
        )

    # -----------------------------------------------------
    # PARSE JSON
    # -----------------------------------------------------

    try:

        result = json.loads(
            response.text
        )

    except json.JSONDecodeError as exc:

        raise RuntimeError(
            "Gemini returned invalid JSON."
        ) from exc

    # -----------------------------------------------------
    # VALIDATE SCORE
    # -----------------------------------------------------

    try:

        result["ats_score"] = max(
            0,
            min(
                100,
                int(result["ats_score"])
            ),
        )

    except (KeyError, TypeError, ValueError) as exc:

        raise RuntimeError(
            "Gemini returned an invalid ATS score."
        ) from exc

    return result


# =========================================================
# DISPLAY RESULTS
# =========================================================

def show_results(result):

    score = result["ats_score"]

    st.subheader("ATS Score")

    st.progress(
        score / 100
    )

    st.metric(
        "Estimated ATS readiness",
        f"{score}/100"
    )

    st.info(
        result["summary"]
    )

    # -----------------------------------------------------
    # STRENGTHS / KEYWORDS
    # -----------------------------------------------------

    col1, col2 = st.columns(2)

    with col1:

        st.subheader(
            "What is already strong"
        )

        for item in result["strengths"]:

            st.markdown(
                f"• {item}"
            )

    with col2:

        st.subheader(
            "Missing / weak keywords"
        )

        if result["keywords_missing"]:

            for item in result["keywords_missing"]:

                st.markdown(
                    f"• `{item}`"
                )

        else:

            st.write(
                "No major missing keywords were identified."
            )

    # -----------------------------------------------------
    # IMPROVEMENTS
    # -----------------------------------------------------

    st.subheader(
        "Recommended improvements"
    )

    for index, item in enumerate(
        result["improvements"],
        start=1,
    ):

        with st.expander(
            f"{index}. {item['issue']} | "
            f"Priority: {item['priority']}"
        ):

            st.write(
                item["recommendation"]
            )

    # -----------------------------------------------------
    # ATS CHECKS
    # -----------------------------------------------------

    st.subheader(
        "ATS checks"
    )

    for check in result["ats_checks"]:

        status = check["status"]

        if status == "Good":

            icon = "✅"

        elif status == "Needs improvement":

            icon = "⚠️"

        else:

            icon = "❌"

        st.markdown(
            f"**{icon} {check['check']}**"
        )

        st.write(
            check["explanation"]
        )


# =========================================================
# MAIN UI
# =========================================================

st.title(
    "📄 Resume ATS Analyzer"
)

st.write(
    "Upload your resume to get an estimated ATS score "
    "and practical improvements."
)


# =========================================================
# SIDEBAR
# =========================================================

with st.sidebar:

    st.header(
        "Settings"
    )

    st.caption(
        f"Gemini model: `{MODEL_NAME}`"
    )

    st.caption(
        "Your resume is processed in memory by this app "
        "and sent to Gemini for analysis. "
        "It is not saved by the app."
    )


# =========================================================
# FILE UPLOAD
# =========================================================

uploaded_file = st.file_uploader(
    "Upload your resume",
    type=[
        "pdf",
        "docx",
        "txt",
    ],
    help=(
        "Supported formats: PDF, DOCX, and TXT."
    ),
)


# =========================================================
# JOB DESCRIPTION
# =========================================================

job_description = st.text_area(
    "Optional: paste the job description",
    height=180,
    placeholder=(
        "Adding the target job description makes "
        "the keyword/ATS analysis more useful."
    ),
)


# =========================================================
# ANALYZE BUTTON
# =========================================================

analyze_button = st.button(
    "Analyze Resume",
    type="primary",
    use_container_width=True,
)


# =========================================================
# ANALYSIS
# =========================================================

if analyze_button:

    # -----------------------------------------------------
    # CHECK FILE
    # -----------------------------------------------------

    if uploaded_file is None:

        st.warning(
            "Please upload a resume first."
        )

        st.stop()

    # -----------------------------------------------------
    # CHECK API KEY
    # -----------------------------------------------------

    api_key = get_api_key()

    if not api_key:

        st.error(
            "Gemini API key is missing. "
            "Add GEMINI_API_KEY to Streamlit Secrets "
            "or set it as an environment variable."
        )

        st.stop()

    # -----------------------------------------------------
    # RUN ANALYSIS
    # -----------------------------------------------------

    try:

        with st.spinner(
            "Reading the resume and generating ATS analysis..."
        ):

            resume_text = extract_resume_text(
                uploaded_file
            )

            # Prevent extremely large documents
            if len(resume_text) > 50000:

                resume_text = resume_text[:50000]

                st.warning(
                    "The resume text was very long, "
                    "so the analysis used the first "
                    "50,000 characters."
                )

            result = analyze_resume(
                resume_text,
                job_description,
                api_key,
            )

        st.success(
            "Analysis complete."
        )

        show_results(
            result
        )

    # -----------------------------------------------------
    # ERROR HANDLING
    # -----------------------------------------------------

    except Exception as exc:

        error_text = str(exc)

        st.error(
            "The analysis could not be completed."
        )

        # Friendly message for 503
        if (
            "503" in error_text
            or "UNAVAILABLE" in error_text
        ):

            st.warning(
                "Gemini is temporarily unavailable "
                "or overloaded. The app automatically "
                "retried the request several times."
            )

            st.caption(
                "Please try the Analyze Resume button again "
                "after a short while."
            )

        # Friendly message for API key problems
        elif (
            "401" in error_text
            or "403" in error_text
            or "API key" in error_text
            or "PERMISSION_DENIED" in error_text
        ):

            st.warning(
                "There appears to be a Gemini API "
                "authentication or permission problem. "
                "Check your GEMINI_API_KEY in Streamlit Secrets."
            )

        else:

            st.exception(
                exc
            )
```

### Ab ek aur important change

Tumhari `requirements.txt` mein **old Google package nahi hona chahiye**.

Usmein ye hona chahiye:

```text
streamlit
google-genai
pypdf
python-docx
```

`google-generativeai` agar requirements mein hai to **remove** kar do. Google ka current SDK `google-genai` hai.

### Streamlit par kya karna hai

1. `app.py` ka pura old code delete karo.
2. Upar wala **complete code paste** karo.
3. Save/commit karo.
4. `requirements.txt` update karo.
5. Streamlit mein **Reboot app** karo.
6. Resume upload karke **Analyze Resume** press karo.

**Ek baat clear:** ye code 503 ko magically eliminate nahi kar sakta. Agar Google ka server genuinely unavailable hai, koi Python change us server ko force nahi kar sakta. Is update ka faida ye hai ke app temporary 503 par **4 attempts with increasing delays** karega instead of immediately crashing. Google ki documentation bhi temporary 503 failures ke liye retry/backoff approach show karti hai.

Agar **is updated code ke baad bhi 503 aaye**, phir next step model change nahi hoga. Phir hum **API key/project side ko isolate karke test** karenge.
